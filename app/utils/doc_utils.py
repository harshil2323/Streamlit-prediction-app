import docx
from docx.shared import Inches, Pt
import streamlit as st

def create_word_document(content, title="Scanned Document", font_name="Calibri", font_size=11):
    """
    Create a Word document from the given content
    """
    try:
        # Create document
        doc = docx.Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Add content
        paragraph = doc.add_paragraph(content)
        
        # Style the paragraph
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
        
        # Set margins (1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        return doc
    except Exception as e:
        st.error(f"Error creating Word document: {str(e)}")
        return None

def save_docx(doc, filename):
    """
    Save the Word document to a file
    """
    try:
        doc.save(filename)
        return True
    except Exception as e:
        st.error(f"Error saving document: {str(e)}")
        return False

def get_document_bytes(doc):
    """
    Get document as bytes for download
    """
    try:
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error converting document to bytes: {str(e)}")
        return None
